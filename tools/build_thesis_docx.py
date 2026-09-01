#!/usr/bin/env python3
"""Build the submission DOCX from the Markdown thesis source.

Renders `manuscript/draft_v5/` (front matter + condensed body + appendices) into a
single .docx laid out to the HS Esslingen *Thesis Writing Guidelines*:

    A4 · 11 pt Arial · 1.15 line spacing · justified · first-line indent
    margins  left 1.5"  right 1"  top/bottom 1.25"
    preliminary pages  lower-case roman, centred at the bottom
    body                arabic, upper right
    body length         80-100 pages  (checked by tools/thesis_page_budget.py)

The Markdown dialect understood here is the one the chapters already use: ATX
headings, GFM tables, `<img src=... width=...>` figures, fenced code, blockquotes,
bullet/numbered lists, `**bold**` / `*italic*` / `` `code` `` inline runs, and
HTML comments (dropped -- they are editorial notes, not thesis text).

Usage:
    python tools/build_thesis_docx.py --out "B:/SE4AI/Documentos/draft_V5.docx"
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parent.parent
SRC = REPO / "manuscript" / "draft_v5"
FIGURES = REPO / "manuscript" / "figures"

# --- guideline constants -------------------------------------------------------
BODY_PT = 12
# Typography, set by the author and deliberately NOT the guidelines' 12 pt Times
# New Roman at 1.5: Arial throughout -- body text, headings and captions alike --
# at 11 pt with 1.15 line spacing. Roman numerals stay confined to the
# preliminary pages (cover, abstract, indices); the body is arabic from 1.
BODY_FONT = "Arial"
BODY_SIZE_PT = 11
LINE_SPACING = 1.15
MARGIN_LEFT = Inches(1.0)
MARGIN_RIGHT = Inches(1.0)
MARGIN_TOPBOT = Inches(1.25)
FIRST_LINE_INDENT = Inches(0.25)  # "indent paragraphs five spaces"
MAX_FIGURE_WIDTH = Inches(5.4)  # text column is 5.77"; leave a hair of slack


# ==============================================================================
# low-level OOXML helpers (python-docx has no API for fields or page numbering)
# ==============================================================================
def _el(tag: str, **attrs) -> OxmlElement:
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    return e


def add_field(paragraph, instr: str, placeholder: str = "") -> None:
    """Insert a Word field (TOC, PAGE, ...) that Word evaluates on open/F9."""
    run = paragraph.add_run()
    run._r.append(_el("w:fldChar", **{"w:fldCharType": "begin"}))
    t = _el("w:instrText", **{"xml:space": "preserve"})
    t.text = instr
    run._r.append(t)
    run._r.append(_el("w:fldChar", **{"w:fldCharType": "separate"}))
    if placeholder:
        ph = _el("w:t")
        ph.text = placeholder
        run._r.append(ph)
    run._r.append(_el("w:fldChar", **{"w:fldCharType": "end"}))


def bookmark(paragraph, name: str) -> None:
    """Wrap a paragraph in a bookmark so a PAGEREF field can point at it."""
    bid = str(abs(hash(name)) % 100000)
    start = _el("w:bookmarkStart", **{"w:id": bid, "w:name": name})
    end = _el("w:bookmarkEnd", **{"w:id": bid})
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def set_page_numbering(section, fmt: str, start: int | None = None) -> None:
    """fmt: 'lowerRoman' | 'decimal'."""
    sectPr = section._sectPr
    for old in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(old)
    attrs = {"w:fmt": fmt}
    if start is not None:
        attrs["w:start"] = str(start)
    sectPr.append(_el("w:pgNumType", **attrs))


def put_page_number(container, alignment) -> None:
    """Drop a bare PAGE field into a header/footer, with no ornamentation."""
    p = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
    p.text = ""
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.first_line_indent = Inches(0)
    add_field(p, " PAGE ", "1")
    for r in p.runs:
        r.font.size = Pt(BODY_PT)


def unlink_headers_footers(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.even_page_header.is_linked_to_previous = False
    section.even_page_footer.is_linked_to_previous = False
    section.first_page_header.is_linked_to_previous = False
    section.first_page_footer.is_linked_to_previous = False


def clear(container) -> None:
    for p in list(container.paragraphs):
        p._element.getparent().remove(p._element)


# ==============================================================================
# document styles
# ==============================================================================
def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE_PT)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = normal.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Make BODY_FONT the document-wide default as well, so anything that never
    # names a font -- page numbers, TOC and index entries, table cells -- resolves
    # to it instead of falling back to the theme font (which Word renders as its
    # own default and LibreOffice substitutes with a serif).
    docdefaults = doc.styles.element.find(qn("w:docDefaults"))
    if docdefaults is not None:
        rpr_default = docdefaults.find(qn("w:rPrDefault"))
        if rpr_default is not None:
            rpr = rpr_default.find(qn("w:rPr"))
            if rpr is None:
                rpr = OxmlElement("w:rPr")
                rpr_default.insert(0, rpr)
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.insert(0, rfonts)
            for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
                if rfonts.get(qn("w:" + attr)) is not None:
                    del rfonts.attrib[qn("w:" + attr)]
            for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                rfonts.set(qn("w:" + attr), BODY_FONT)

    # Subheadings: flush left, same family, distinguished only by weight/size.
    # "three or four blank lines before, two after" -> rendered as space_before/after.
    for name, size, before, after, bold, caps in (
        ("Heading 1", 16, 0, 24, True, True),
        ("Heading 2", 14, 24, 12, True, False),
        ("Heading 3", 12, 18, 10, True, False),
        ("Heading 4", 12, 14, 8, False, False),
    ):
        st = doc.styles[name]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.all_caps = caps
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.font.italic = False
        # python-docx's heading styles point at the theme's *major* font, which wins
        # over w:ascii; drop the theme references so BODY_FONT actually applies.
        rfonts = st.element.rPr.rFonts
        for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
            if rfonts.get(qn("w:" + attr)) is not None:
                del rfonts.attrib[qn("w:" + attr)]
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn("w:" + attr), BODY_FONT)
        p = st.paragraph_format
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.space_before = Pt(before)
        p.space_after = Pt(after)
        p.line_spacing = 1.15
        p.keep_with_next = True
        p.first_line_indent = Inches(0)

    for name in ("Caption", "Quote", "List Bullet", "List Number"):
        if name in [s.name for s in doc.styles]:
            st = doc.styles[name]
            st.font.name = BODY_FONT


def force_theme_font(doc: Document, font: str) -> None:
    """Point the theme's major/minor latin fonts at `font`.

    Styles that python-docx's default template leaves on `majorHAnsi`/`minorHAnsi`
    -- table styles, TOC and index entries, headers and footers -- resolve through
    the theme, which ships Calibri/Cambria. Without this the document is Arial
    everywhere except a handful of places that quietly come out serif.
    """
    for part in doc.part.package.iter_parts():
        if not str(part.partname).endswith("theme1.xml"):
            continue
        xml = part.blob.decode("utf-8")
        xml = re.sub(r'(<a:latin typeface=")[^"]*(")', rf"\1{font}\2", xml)
        part._blob = xml.encode("utf-8")


def set_margins(section) -> None:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.top_margin = MARGIN_TOPBOT
    section.bottom_margin = MARGIN_TOPBOT
    section.header_distance = Inches(0.6)
    section.footer_distance = Inches(0.6)


# ==============================================================================
# inline markdown -> runs
# ==============================================================================
CAPTION_RE = re.compile(r"^(Figura|Tabla)\s+[\d.]+\s*[—–-]")

INLINE = re.compile(
    r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|(?<!\*)\*(?!\s)[^*]+?\*|`[^`]+`|\[[^\]]+\]\([^)]+\))",
    re.S,
)


def add_inline(paragraph, text: str, base_italic: bool = False) -> None:
    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        return
    for chunk in INLINE.split(text):
        if not chunk:
            continue
        bold = italic = mono = False
        if chunk.startswith("***") and chunk.endswith("***"):
            chunk, bold, italic = chunk[3:-3], True, True
        elif chunk.startswith("**") and chunk.endswith("**"):
            chunk, bold = chunk[2:-2], True
        elif chunk.startswith("*") and chunk.endswith("*"):
            chunk, italic = chunk[1:-1], True
        elif chunk.startswith("`") and chunk.endswith("`"):
            chunk, mono = chunk[1:-1], True
        else:
            m = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", chunk)
            if m:
                chunk = m.group(1)
        chunk = (
            chunk.replace("\\|", "|")
            .replace("\\*", "*")
            .replace("\\_", "_")
            .replace("&nbsp;", " ")
        )
        run = paragraph.add_run(chunk)
        run.bold = bold
        run.italic = italic or base_italic
        if mono:
            run.font.name = "Consolas"
            run.font.size = Pt(BODY_PT - 1.5)


# ==============================================================================
# block rendering
# ==============================================================================
class Builder:
    def __init__(self, doc: Document):
        self.doc = doc
        self.figure_count = 0
        self.missing_figures: list[str] = []
        self.line_spacing = LINE_SPACING  # the abstract is single-spaced (guidelines)
        self.centred = False  # the title page is centred and unindented
        self.captions: list[tuple[str, str]] = []  # (visible label, bookmark name)

    # -- paragraphs ------------------------------------------------------------
    def para(self, text, *, style=None, indent=True, italic=False, align=None):
        p = self.doc.add_paragraph(style=style)
        if self.centred:
            align, indent = WD_ALIGN_PARAGRAPH.CENTER, False
        if align is not None:
            p.alignment = align
        if indent:
            p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
        else:
            p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.line_spacing = self.line_spacing
        p.paragraph_format.space_after = Pt(0)
        add_inline(p, text, base_italic=italic)
        if CAPTION_RE.match(text.strip().lstrip('*')):
            self.register_caption(p, text)
        return p

    def register_caption(self, paragraph, text: str) -> None:
        label = re.sub(r"[*_`]", "", text).strip()
        label = re.sub(r"\s+", " ", label)
        if len(label) > 118:
            label = label[:115].rstrip(" ,;:") + "..."
        name = f"_cap{len(self.captions):03d}"
        bookmark(paragraph, name)
        self.captions.append((label, name))

    def heading(self, text, level):
        p = self.doc.add_paragraph(style=f"Heading {min(level, 4)}")
        add_inline(p, text)
        return p

    def code_block(self, lines):
        for ln in lines:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf = p.paragraph_format
            pf.line_spacing = 1.0
            pf.space_after = Pt(0)
            pf.left_indent = Inches(0.3)
            pf.first_line_indent = Inches(0)
            r = p.add_run(ln if ln.strip() else " ")
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def quote(self, text):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Inches(0.4)
        pf.right_indent = Inches(0.4)
        pf.first_line_indent = Inches(0)
        pf.space_after = Pt(6)
        add_inline(p, text)

    def bullet(self, text, style="List Bullet", level=0):
        p = self.doc.add_paragraph(style=style)
        pf = p.paragraph_format
        pf.left_indent = Inches(0.35 + 0.3 * level)
        pf.first_line_indent = Inches(0)
        pf.space_after = Pt(0)
        add_inline(p, text)

    # -- figures ---------------------------------------------------------------
    def figure(self, src, width_px=None, caption=None):
        name = Path(src).name
        for candidate in (FIGURES / name, FIGURES / "auto" / name):
            if candidate.exists():
                path = candidate.resolve()
                break
        else:
            self.missing_figures.append(src)
            return
        width = MAX_FIGURE_WIDTH
        if width_px:
            requested = Inches(int(width_px) / 96)
            width = min(requested, MAX_FIGURE_WIDTH)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        try:
            p.add_run().add_picture(str(path), width=width)
        except Exception as exc:  # pragma: no cover - malformed asset
            self.missing_figures.append(f"{src} ({exc})")
            return
        self.figure_count += 1
        if caption:
            c = self.doc.add_paragraph()
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraph_format.first_line_indent = Inches(0)
            c.paragraph_format.space_after = Pt(10)
            add_inline(c, caption, base_italic=True)
            for r in c.runs:
                r.font.size = Pt(10)
            self.register_caption(c, caption)

    # -- tables ----------------------------------------------------------------
    def table(self, rows):
        header, body = rows[0], rows[1:]
        t = self.doc.add_table(rows=len(rows), cols=len(header))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = True
        for j, cell in enumerate(header):
            self._cell(t.cell(0, j), cell, bold=True)
        for i, row in enumerate(body, start=1):
            for j in range(len(header)):
                self._cell(t.cell(i, j), row[j] if j < len(row) else "")
        after = self.doc.add_paragraph()
        after.paragraph_format.space_after = Pt(8)
        after.paragraph_format.line_spacing = 1.0

    @staticmethod
    def _cell(cell, text, bold=False):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.line_spacing = 1.0
        pf.space_after = Pt(2)
        pf.space_before = Pt(2)
        pf.first_line_indent = Inches(0)
        add_inline(p, text)
        for r in p.runs:
            r.font.size = Pt(9.5)
            if bold:
                r.bold = True

    def page_break(self):
        self.doc.add_page_break()


# ==============================================================================
# markdown driver
# ==============================================================================
IMG_RE = re.compile(r'<img\s+src="([^"]+)"[^>]*?(?:width="(\d+)")?[^>]*/?>', re.I)
ROW_SPLIT = re.compile(r"(?<!\\)\|")


def split_row(line):
    cells = ROW_SPLIT.split(line.strip())
    if cells and not cells[0].strip():
        cells = cells[1:]
    if cells and not cells[-1].strip():
        cells = cells[:-1]
    return [c.strip() for c in cells]


def render_markdown(builder: Builder, text: str, heading_offset: int = 0) -> None:
    text = re.sub(r"<!--.*?-->", "", text, flags=re.S)
    builder.line_spacing = 1.0 if "::single-space::" in text else LINE_SPACING
    builder.centred = "::centre::" in text
    text = text.replace("::single-space::", "").replace("::centre::", "")
    lines = text.split("\n")
    i, n = 0, len(lines)
    para_buf: list[str] = []

    def flush(italic=False):
        nonlocal para_buf
        if para_buf:
            builder.para(" ".join(para_buf), italic=italic)
            para_buf = []

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            flush()
            i += 1
            continue

        # fenced code / mermaid
        if stripped.startswith("```"):
            flush()
            lang = stripped[3:].strip().lower()
            j = i + 1
            block = []
            while j < n and not lines[j].strip().startswith("```"):
                block.append(lines[j])
                j += 1
            if lang == "mermaid":
                # prefer a rendered PNG named in the "Fuente canónica" comment
                src = None
                for b in block:
                    m = re.search(r"([\w./-]+\.mmd)", b)
                    if m:
                        src = Path(m.group(1)).with_suffix(".png").name
                        break
                if src and (FIGURES / src).exists():
                    builder.figure(src)
            else:
                builder.code_block(block)
            i = j + 1
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            flush()
            level = len(m.group(1)) + heading_offset
            title = m.group(2).strip()
            title = re.sub(r"\s*\[[^\]]*\]\s*$", "", title)  # drop [BORRADOR ...] tags
            builder.heading(title, level)
            i += 1
            continue

        # horizontal rule
        if re.fullmatch(r"[-*_]{3,}", stripped):
            flush()
            i += 1
            continue

        # image
        mi = IMG_RE.search(stripped)
        if mi:
            flush()
            caption = None
            k = i + 1
            while k < n and not lines[k].strip():
                k += 1
            if k < n and re.match(r"^\**\*?(Figura|Fig\.)", lines[k].strip()):
                caption = lines[k].strip().strip("*")
                k += 1
            builder.figure(mi.group(1), mi.group(2), caption)
            i = k
            continue

        # table
        if stripped.startswith("|"):
            flush()
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                r = split_row(lines[i])
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in r if c):
                    rows.append(r)
                i += 1
            if rows:
                builder.table(rows)
            continue

        # blockquote
        if stripped.startswith(">"):
            flush()
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip().lstrip(">").strip())
                i += 1
            builder.quote(" ".join(x for x in buf if x))
            continue

        # lists
        ml = re.match(r"^(\s*)([-*+]|\d+[.)])\s+(.*)$", line)
        if ml:
            flush()
            indent = len(ml.group(1))
            ordered = not ml.group(2) in "-*+"
            item = [ml.group(3)]
            i += 1
            while i < n:
                nxt = lines[i]
                if not nxt.strip():
                    break
                if re.match(r"^(\s*)([-*+]|\d+[.)])\s+", nxt):
                    break
                if nxt.startswith("|") or nxt.strip().startswith("#"):
                    break
                item.append(nxt.strip())
                i += 1
            builder.bullet(
                " ".join(item),
                style="List Number" if ordered else "List Bullet",
                level=min(indent // 2, 2),
            )
            continue

        para_buf.append(stripped)
        i += 1

    flush()


# ==============================================================================
# assembly
# ==============================================================================
def read(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def build(out_path: Path) -> None:
    doc = Document()
    configure_styles(doc)
    force_theme_font(doc, BODY_FONT)

    front_sec = doc.sections[0]
    set_margins(front_sec)
    unlink_headers_footers(front_sec)
    builder = Builder(doc)

    # ---- cover page (no number) ------------------------------------------
    cover = SRC / "front" / "00_cover.md"
    render_markdown(builder, read(cover))
    builder.page_break()

    # ---- preliminary pages: roman numerals, centred at the bottom ---------
    for name in sorted((SRC / "front").glob("*.md")):
        if name.name == "00_cover.md":
            continue
        render_markdown(builder, read(name))
        if name.name == "20_toc.md":
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0)
            add_field(p, r' TOC \o "1-3" \h \z \u ', "Actualice el índice con F9.")
        if name.name == "30_list_of_figures.md":
            lof_anchor = doc.paragraphs[-1]
        builder.page_break()

    set_page_numbering(front_sec, "lowerRoman", start=1)
    clear(front_sec.footer)
    put_page_number(front_sec.footer, WD_ALIGN_PARAGRAPH.CENTER)
    clear(front_sec.first_page_footer)  # title page carries no number
    front_sec.different_first_page_header_footer = True

    # ---- body: arabic, upper right, restarting at 1 -----------------------
    body_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(body_sec)
    unlink_headers_footers(body_sec)
    set_page_numbering(body_sec, "decimal", start=1)
    clear(body_sec.footer)
    clear(body_sec.header)
    put_page_number(body_sec.header, WD_ALIGN_PARAGRAPH.RIGHT)
    body_sec.different_first_page_header_footer = False

    body_files = sorted((SRC / "body").glob("*.md"))
    for k, f in enumerate(body_files):
        if k:
            builder.page_break()
        render_markdown(builder, read(f))

    body_end_marker = len(doc.paragraphs)

    # ---- back matter: bibliography + appendices ---------------------------
    for f in sorted((SRC / "back").glob("*.md")):
        builder.page_break()
        render_markdown(builder, read(f))

    # ---- fill the list of figures and tables (PAGEREF resolves on field update) --
    anchor = lof_anchor._p
    for label, mark in builder.captions:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.first_line_indent = Inches(0)
        pf.left_indent = Inches(0.3)
        pf.first_line_indent = Inches(-0.3)
        pf.line_spacing = 1.0
        pf.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(label + "  ")
        run.font.size = Pt(11)
        add_field(p, rf" PAGEREF {mark} \h ", "-")
        for r in p.runs:
            r.font.size = Pt(11)
        anchor.addnext(p._p)
        anchor = p._p

    doc.save(out_path)
    print(f"written: {out_path}")
    print(f"captions listed: {len(builder.captions)}")
    print(f"figures embedded: {builder.figure_count}")
    if builder.missing_figures:
        print("MISSING FIGURES:")
        for s in sorted(set(builder.missing_figures)):
            print("   -", s)
    print(f"paragraphs at end of body: {body_end_marker}")


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument(
        "--out",
        default=str(Path("B:/SE4AI/Documentos/draft_V5.docx")),
        help="output .docx path",
    )
    args = ap.parse_args()
    if not SRC.exists():
        print(f"source tree not found: {SRC}", file=sys.stderr)
        return 1
    build(Path(args.out))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
